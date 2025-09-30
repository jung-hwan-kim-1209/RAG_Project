"""
Layer 8: REPORT GENERATION LAYER
unicorn_report_generator를 실행하여 최종 투자 평가 리포트를 생성하는 레이어
"""
import os
from typing import List, Dict, Any, Optional
from datetime import datetime
from langchain_openai import ChatOpenAI
from langchain_core.prompts import PromptTemplate
from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import platform

import sys
import os
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from models import (
    InvestmentReport, InvestmentRecommendation, UnicornScore,
    AnalysisResult, RiskAssessment, CompanyInfo, PipelineContext, RiskLevel,
    DocumentChunk, ExternalSearchResult, GPTResponse
)
from config import get_config

class UnicornReportGenerator:
    """유니콘 투자 평가 리포트 생성기"""

    def __init__(self):
        self.config = get_config()
        self.llm = ChatOpenAI(
            openai_api_key=self.config["model"].openai_api_key,
            temperature=0.1,
            model=self.config["model"].model_name  # 예: "gpt-4o-mini"
        )

        # Executive Summary 생성 프롬프트
        self.executive_summary_prompt = PromptTemplate(
            input_variables=["company_name", "total_score", "grade", "unicorn_probability", "recommendation"],
            template="""다음 정보를 바탕으로 {company_name}에 대한 투자 평가 Executive Summary를 작성해주세요.

회사명: {company_name}
종합 점수: {total_score}/100
등급: {grade}
유니콘 확률: {unicorn_probability:.1%}
투자 추천: {recommendation}

Executive Summary는 다음 구조로 작성해주세요:
1. 핵심 결론 (2-3줄)
2. 주요 강점 (3-4개 항목)
3. 주요 우려사항 (2-3개 항목)
4. 투자 권장사항 (1-2줄)

전문적이고 간결한 투자 보고서 형식으로 작성해주세요."""
        )

        # 상세 분석 요약 프롬프트
        self.detailed_analysis_prompt = PromptTemplate(
            input_variables=["company_name", "analysis_results", "risk_assessments"],
            template="""다음 분석 결과들을 바탕으로 {company_name}에 대한 상세 분석 보고서를 작성해주세요.

분석 결과:
{analysis_results}

리스크 평가:
{risk_assessments}

다음 섹션별로 구성해주세요:
1. 성장성 분석
2. 비즈니스 모델 평가
3. 기술력 및 보안성
4. 재무 건전성

각 섹션은 점수, 주요 발견사항, 개선 권장사항을 포함해주세요."""
        )

        # 투자 근거 프롬프트
        self.investment_rationale_prompt = PromptTemplate(
            input_variables=["company_name", "recommendation", "unicorn_score", "key_factors"],
            template="""다음 정보를 바탕으로 {company_name}에 대한 투자 근거를 작성해주세요.

투자 추천: {recommendation}
유니콘 점수: {unicorn_score}
주요 요인들: {key_factors}

투자 근거는 다음을 포함해주세요:
1. 투자 결정의 핵심 논리
2. 예상 수익률 및 리스크 균형
3. 투자 타이밍의 적절성
4. 포트폴리오 내 포지셔닝
5. Exit 전략 고려사항

투자자 관점에서 명확하고 설득력 있게 작성해주세요."""
        )

    def determine_investment_recommendation(
        self,
        unicorn_score: UnicornScore,
        risk_assessments: List[RiskAssessment]
    ) -> InvestmentRecommendation:
        """투자 추천 결정"""

        total_score = unicorn_score.total_score
        unicorn_probability = unicorn_score.unicorn_probability

        # 높은 리스크 체크
        high_risk_count = sum(1 for risk in risk_assessments
                             if risk.risk_level.value in ["높음", "매우 높음"])

        # 투자 결정 로직
        if total_score >= 85 and unicorn_probability >= 0.7 and high_risk_count <= 1:
            return InvestmentRecommendation.INVEST
        elif total_score >= 70 and unicorn_probability >= 0.5 and high_risk_count <= 2:
            return InvestmentRecommendation.INVEST
        elif total_score >= 60 and high_risk_count <= 3:
            return InvestmentRecommendation.HOLD
        else:
            return InvestmentRecommendation.AVOID

    def generate_executive_summary(
        self,
        company_info: CompanyInfo,
        unicorn_score: UnicornScore,
        recommendation: InvestmentRecommendation
    ) -> str:
        """Executive Summary 생성"""
        try:
            response = self.llm.invoke(self.executive_summary_prompt.format(
                company_name=company_info.name,
                total_score=unicorn_score.total_score,
                grade=unicorn_score.grade,
                unicorn_probability=unicorn_score.unicorn_probability,
                recommendation=recommendation.value
            ))

            # GPT 응답을 터미널에 출력
            print(f"\n🔍 REPORT_GENERATION_LAYER (EXECUTIVE_SUMMARY) - GPT 응답:")
            print("=" * 60)
            print(response.content)
            print("=" * 60)

            return response.content.strip()
        except Exception as e:
            return f"Executive Summary 생성 오류: {str(e)}"

    def generate_detailed_analysis(
        self,
        company_info: CompanyInfo,
        analysis_results: List[AnalysisResult],
        risk_assessments: List[RiskAssessment]
    ) -> str:
        """상세 분석 보고서 생성"""
        try:
            # 분석 결과 포맷팅
            analysis_text = []
            for result in analysis_results:
                analysis_text.append(
                    f"[{result.category}] {result.score}점 ({result.grade}급)\n"
                    f"요약: {result.summary}\n"
                    f"강점: {', '.join(result.key_strengths)}\n"
                    f"약점: {', '.join(result.key_weaknesses)}\n"
                )

            # 리스크 평가 포맷팅
            risk_text = []
            for risk in risk_assessments:
                risk_text.append(
                    f"[{risk.category}] {risk.risk_level.value}\n"
                    f"설명: {risk.description}\n"
                    f"완화 전략: {', '.join(risk.mitigation_strategies)}\n"
                )

            response = self.llm.invoke(self.detailed_analysis_prompt.format(
                company_name=company_info.name,
                analysis_results="\n".join(analysis_text),
                risk_assessments="\n".join(risk_text)
            ))

            # GPT 응답을 터미널에 출력
            print(f"\n🔍 REPORT_GENERATION_LAYER (DETAILED_ANALYSIS) - GPT 응답:")
            print("=" * 60)
            print(response.content)
            print("=" * 60)

            return response.content.strip()
        except Exception as e:
            return f"상세 분석 생성 오류: {str(e)}"

    def generate_investment_rationale(
        self,
        company_info: CompanyInfo,
        recommendation: InvestmentRecommendation,
        unicorn_score: UnicornScore
    ) -> str:
        """투자 근거 생성"""
        try:
            key_factors = unicorn_score.score_breakdown.get("unicorn_factors", [])
            key_factors_text = ", ".join(key_factors) if key_factors else "종합 분석 결과"

            response = self.llm.invoke(self.investment_rationale_prompt.format(
                company_name=company_info.name,
                recommendation=recommendation.value,
                unicorn_score=f"{unicorn_score.total_score:.1f}점 ({unicorn_score.grade}급)",
                key_factors=key_factors_text
            ))

            # GPT 응답을 터미널에 출력
            print(f"\n🔍 REPORT_GENERATION_LAYER (INVESTMENT_RATIONALE) - GPT 응답:")
            print("=" * 60)
            print(response.content)
            print("=" * 60)

            return response.content.strip()
        except Exception as e:
            return f"투자 근거 생성 오류: {str(e)}"

    def generate_risk_summary(self, risk_assessments: List[RiskAssessment]) -> str:
        """리스크 요약 생성"""
        if not risk_assessments:
            return "리스크 평가 결과가 없습니다."

        risk_summary = []
        risk_summary.append("## 리스크 요약")

        # 리스크 레벨별 분류
        critical_risks = [r for r in risk_assessments if r.risk_level == RiskLevel.CRITICAL]
        high_risks = [r for r in risk_assessments if r.risk_level == RiskLevel.HIGH]
        medium_risks = [r for r in risk_assessments if r.risk_level == RiskLevel.MEDIUM]
        low_risks = [r for r in risk_assessments if r.risk_level == RiskLevel.LOW]

        if critical_risks:
            risk_summary.append("### 🔴 매우 높은 리스크")
            for risk in critical_risks:
                risk_summary.append(f"- **{risk.category}**: {risk.description}")

        if high_risks:
            risk_summary.append("### 🟠 높은 리스크")
            for risk in high_risks:
                risk_summary.append(f"- **{risk.category}**: {risk.description}")

        if medium_risks:
            risk_summary.append("### 🟡 보통 리스크")
            for risk in medium_risks:
                risk_summary.append(f"- **{risk.category}**: {risk.description}")

        if low_risks:
            risk_summary.append("### 🟢 낮은 리스크")
            for risk in low_risks:
                risk_summary.append(f"- **{risk.category}**: {risk.description}")

        return "\n".join(risk_summary)

    def calculate_confidence_level(
        self,
        analysis_results: List[AnalysisResult],
        risk_assessments: List[RiskAssessment],
        documents: List[DocumentChunk],
        external_results: List[ExternalSearchResult]
    ) -> float:
        """신뢰도 레벨 계산"""
        confidence_factors = []

        # 데이터 양 기반 신뢰도
        max_data_sources = int(os.getenv("MAX_DATA_SOURCES_FOR_CONFIDENCE", "20"))
        data_confidence = min(len(documents) + len(external_results), max_data_sources) / max_data_sources
        confidence_factors.append(data_confidence)

        # 분석 완성도 기반 신뢰도
        max_analysis_areas = int(os.getenv("MAX_ANALYSIS_AREAS", "4"))
        analysis_confidence = len(analysis_results) / max_analysis_areas
        confidence_factors.append(analysis_confidence)

        # 리스크 평가 완성도
        max_risk_categories = int(os.getenv("MAX_RISK_CATEGORIES", "6"))
        risk_confidence = len(risk_assessments) / max_risk_categories
        confidence_factors.append(risk_confidence)

        # 평균 신뢰도 계산
        overall_confidence = sum(confidence_factors) / len(confidence_factors)

        return min(overall_confidence, 1.0)

    def generate_investment_report(
        self,
        company_info: CompanyInfo,
        unicorn_score: UnicornScore,
        analysis_results: List[AnalysisResult],
        risk_assessments: List[RiskAssessment],
        documents: List[DocumentChunk],
        external_results: List[ExternalSearchResult]
    ) -> InvestmentReport:
        """최종 투자 평가 리포트 생성"""

        # 투자 추천 결정
        recommendation = self.determine_investment_recommendation(unicorn_score, risk_assessments)

        # 각 섹션별 내용 생성
        executive_summary = self.generate_executive_summary(
            company_info, unicorn_score, recommendation
        )

        detailed_analysis = self.generate_detailed_analysis(
            company_info, analysis_results, risk_assessments
        )

        investment_rationale = self.generate_investment_rationale(
            company_info, recommendation, unicorn_score
        )

        risk_summary = self.generate_risk_summary(risk_assessments)

        # 신뢰도 계산
        confidence_level = self.calculate_confidence_level(
            analysis_results, risk_assessments, documents, external_results
        )

        # 데이터 소스 정리
        data_sources = []
        for doc in documents:
            data_sources.append(doc.source)
        for result in external_results:
            data_sources.append(result.source)
        data_sources = list(set(data_sources))  # 중복 제거

        # 제한사항 정리
        limitations = []
        min_documents = int(os.getenv("MIN_DOCUMENTS_FOR_LIMITATION", "5"))
        min_external_results = int(os.getenv("MIN_EXTERNAL_RESULTS_FOR_LIMITATION", "3"))
        min_confidence_level = float(os.getenv("MIN_CONFIDENCE_LEVEL_FOR_LIMITATION", "0.7"))
        
        if len(documents) < min_documents:
            limitations.append("제한된 내부 문서 데이터")
        if len(external_results) < min_external_results:
            limitations.append("제한된 외부 검색 결과")
        if confidence_level < min_confidence_level:
            limitations.append("중간 수준의 분석 신뢰도")

        # 최종 리포트 생성
        report = InvestmentReport(
            company_info=company_info,
            evaluation_date=datetime.now(),
            unicorn_score=unicorn_score,
            recommendation=recommendation,
            analysis_results=analysis_results,
            risk_assessments=risk_assessments,
            executive_summary=executive_summary,
            detailed_analysis=detailed_analysis,
            investment_rationale=investment_rationale,
            risk_summary=risk_summary,
            confidence_level=confidence_level,
            data_sources=data_sources,
            limitations=limitations
        )

        return report

class ReportFormatter:
    """리포트 포맷터"""

    def __init__(self):
        self._register_korean_fonts()

    def _register_korean_fonts(self):
        """한국어 폰트 등록"""
        try:
            # Windows 시스템에서 기본 한국어 폰트 경로들
            korean_fonts = [
                # Windows 기본 폰트
                ("C:/Windows/Fonts/malgun.ttf", "MalgunGothic"),
                ("C:/Windows/Fonts/gulim.ttc", "Gulim"),
                ("C:/Windows/Fonts/batang.ttc", "Batang"),
                ("C:/Windows/Fonts/dotum.ttc", "Dotum"),
                ("C:/Windows/Fonts/arial.ttf", "Arial"),
                ("C:/Windows/Fonts/calibri.ttf", "Calibri"),
                # macOS 기본 폰트
                ("/System/Library/Fonts/AppleGothic.ttf", "AppleGothic"),
                ("/System/Library/Fonts/Helvetica.ttc", "Helvetica"),
                # Linux 기본 폰트
                ("/usr/share/fonts/truetype/nanum/NanumGothic.ttf", "NanumGothic"),
                ("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", "DejaVuSans"),
            ]
            
            # 시스템에 따라 폰트 등록 시도
            for font_path, font_name in korean_fonts:
                try:
                    if os.path.exists(font_path):
                        pdfmetrics.registerFont(TTFont(font_name, font_path))
                        print(f"✅ 한국어 폰트 등록 성공: {font_name}")
                        return font_name
                except Exception as e:
                    print(f"⚠️ 폰트 등록 실패 {font_path}: {e}")
                    continue
            
            # 폰트 등록 실패 시 기본 폰트 사용
            print("⚠️ 한국어 폰트를 찾을 수 없습니다. 기본 폰트를 사용합니다.")
            return "Helvetica"
            
        except Exception as e:
            print(f"⚠️ 폰트 등록 중 오류: {e}")
            return "Helvetica"

    def _get_korean_font_name(self):
        """Word용 한국어 폰트 이름 반환"""
        try:
            # Windows 시스템에서 기본 한국어 폰트들
            korean_fonts = [
                "Malgun Gothic",  # Windows 10/11 기본
                "Gulim",          # Windows 기본
                "Batang",         # Windows 기본
                "Dotum",          # Windows 기본
                "AppleGothic",    # macOS
                "NanumGothic",    # Linux
                "Arial Unicode MS" # 범용
            ]
            
            # 첫 번째로 사용 가능한 폰트 반환
            return korean_fonts[0]
            
        except Exception as e:
            print(f"⚠️ Word 폰트 설정 중 오류: {e}")
            return "Arial"

    def format_gpt_responses_section(self, gpt_responses: List[GPTResponse]) -> str:
        """GPT 응답들을 리포트 섹션으로 포맷팅"""
        if not gpt_responses:
            return "GPT 응답 로그가 없습니다."
        
        formatted_section = "\n" + "="*80 + "\n"
        formatted_section += "🤖 GPT 분석 로그\n"
        formatted_section += "="*80 + "\n\n"
        
        # 레이어별로 그룹화
        layer_groups = {}
        for response in gpt_responses:
            if response.layer_name not in layer_groups:
                layer_groups[response.layer_name] = []
            layer_groups[response.layer_name].append(response)
        
        for layer_name, responses in layer_groups.items():
            formatted_section += f"📋 {layer_name}\n"
            formatted_section += "-" * 60 + "\n"
            
            for i, response in enumerate(responses, 1):
                formatted_section += f"\n{i}. {response.analyzer_name.upper()}\n"
                formatted_section += f"   시간: {response.timestamp.strftime('%Y-%m-%d %H:%M:%S')}\n"
                formatted_section += f"   프롬프트: {response.prompt[:200]}...\n"
                formatted_section += f"   응답: {response.response[:500]}...\n"
                formatted_section += "\n"
        
        return formatted_section

    def export_to_pdf(self, report: InvestmentReport, gpt_responses: List[GPTResponse], output_path: str) -> bool:
        """PDF 파일로 내보내기"""
        try:
            # 한국어 폰트 등록
            korean_font = self._register_korean_fonts()
            
            doc = SimpleDocTemplate(output_path, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            
            # 한국어 폰트를 사용하는 스타일 생성
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontName=korean_font,
                fontSize=18,
                spaceAfter=30,
                alignment=1  # 중앙 정렬
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontName=korean_font,
                fontSize=14,
                spaceAfter=12
            )
            
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontName=korean_font,
                fontSize=10,
                spaceAfter=6
            )
            
            story.append(Paragraph("🦄 AI 스타트업 투자 평가 리포트", title_style))
            story.append(Spacer(1, 12))
            
            # 기본 정보
            story.append(Paragraph("📊 기본 정보", heading_style))
            basic_info = [
                ['회사명', report.company_info.name],
                ['업종', report.company_info.industry],
                ['평가일', report.evaluation_date.strftime('%Y-%m-%d')],
                ['총점', f"{report.unicorn_score.total_score}/100"],
                ['등급', report.unicorn_score.grade],
                ['유니콘 확률', f"{report.unicorn_score.unicorn_probability:.1%}"],
                ['투자 추천', report.recommendation.value]
            ]
            
            basic_table = Table(basic_info, colWidths=[2*inch, 4*inch])
            basic_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
                ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, -1), korean_font),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ]))
            story.append(basic_table)
            story.append(Spacer(1, 20))
            
            # Executive Summary
            story.append(Paragraph("📝 Executive Summary", heading_style))
            story.append(Paragraph(report.executive_summary, normal_style))
            story.append(Spacer(1, 20))
            
            # 상세 분석
            story.append(Paragraph("📈 상세 분석", heading_style))
            story.append(Paragraph(report.detailed_analysis, normal_style))
            story.append(Spacer(1, 20))
            
            # 투자 근거
            story.append(Paragraph("💡 투자 근거", heading_style))
            story.append(Paragraph(report.investment_rationale, normal_style))
            story.append(Spacer(1, 20))
            
            # 리스크 요약
            story.append(Paragraph("⚠️ 리스크 요약", heading_style))
            story.append(Paragraph(report.risk_summary, normal_style))
            story.append(Spacer(1, 20))
            
            # GPT 응답 로그
            if gpt_responses:
                story.append(Paragraph("🤖 GPT 분석 로그", heading_style))
                
                for response in gpt_responses:
                    gpt_heading_style = ParagraphStyle(
                        'GPTHeading',
                        parent=styles['Heading3'],
                        fontName=korean_font,
                        fontSize=12,
                        spaceAfter=6
                    )
                    story.append(Paragraph(f"<b>{response.layer_name} - {response.analyzer_name}</b>", gpt_heading_style))
                    story.append(Paragraph(f"시간: {response.timestamp.strftime('%Y-%m-%d %H:%M:%S')}", normal_style))
                    story.append(Paragraph(f"프롬프트: {response.prompt[:300]}...", normal_style))
                    story.append(Paragraph(f"응답: {response.response[:500]}...", normal_style))
                    story.append(Spacer(1, 12))
            
            # PDF 생성
            doc.build(story)
            return True
            
        except Exception as e:
            print(f"PDF 내보내기 실패: {e}")
            return False

    def export_to_word(self, report: InvestmentReport, gpt_responses: List[GPTResponse], output_path: str) -> bool:
        """Word 파일로 내보내기"""
        try:
            doc = Document()
            
            # 한국어 폰트 설정
            korean_font = self._get_korean_font_name()
            
            # 제목
            title = doc.add_heading('🦄 AI 스타트업 투자 평가 리포트', 0)
            title.alignment = 1  # 중앙 정렬
            
            # 제목 폰트 설정
            for run in title.runs:
                run.font.name = korean_font
                run.font.size = Inches(0.2)
            
            # 기본 정보
            heading1 = doc.add_heading('📊 기본 정보', level=1)
            for run in heading1.runs:
                run.font.name = korean_font
                
            basic_info = [
                ['회사명', report.company_info.name],
                ['업종', report.company_info.industry],
                ['평가일', report.evaluation_date.strftime('%Y-%m-%d')],
                ['총점', f"{report.unicorn_score.total_score}/100"],
                ['등급', report.unicorn_score.grade],
                ['유니콘 확률', f"{report.unicorn_score.unicorn_probability:.1%}"],
                ['투자 추천', report.recommendation.value]
            ]
            
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '항목'
            hdr_cells[1].text = '값'
            
            # 테이블 폰트 설정
            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = korean_font
            
            for item, value in basic_info:
                row_cells = table.add_row().cells
                row_cells[0].text = item
                row_cells[1].text = value
                
                # 각 셀의 폰트 설정
                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = korean_font
            
            # Executive Summary
            heading2 = doc.add_heading('📝 Executive Summary', level=1)
            for run in heading2.runs:
                run.font.name = korean_font
            para1 = doc.add_paragraph(report.executive_summary)
            for run in para1.runs:
                run.font.name = korean_font
            
            # 상세 분석
            heading3 = doc.add_heading('📈 상세 분석', level=1)
            for run in heading3.runs:
                run.font.name = korean_font
            para2 = doc.add_paragraph(report.detailed_analysis)
            for run in para2.runs:
                run.font.name = korean_font
            
            # 투자 근거
            heading4 = doc.add_heading('💡 투자 근거', level=1)
            for run in heading4.runs:
                run.font.name = korean_font
            para3 = doc.add_paragraph(report.investment_rationale)
            for run in para3.runs:
                run.font.name = korean_font
            
            # 리스크 요약
            heading5 = doc.add_heading('⚠️ 리스크 요약', level=1)
            for run in heading5.runs:
                run.font.name = korean_font
            para4 = doc.add_paragraph(report.risk_summary)
            for run in para4.runs:
                run.font.name = korean_font
            
            # GPT 응답 로그
            if gpt_responses:
                gpt_heading = doc.add_heading('🤖 GPT 분석 로그', level=1)
                for run in gpt_heading.runs:
                    run.font.name = korean_font
                
                for response in gpt_responses:
                    sub_heading = doc.add_heading(f"{response.layer_name} - {response.analyzer_name}", level=2)
                    for run in sub_heading.runs:
                        run.font.name = korean_font
                        
                    para_time = doc.add_paragraph(f"시간: {response.timestamp.strftime('%Y-%m-%d %H:%M:%S')}")
                    for run in para_time.runs:
                        run.font.name = korean_font
                        
                    para_prompt = doc.add_paragraph(f"프롬프트: {response.prompt[:300]}...")
                    for run in para_prompt.runs:
                        run.font.name = korean_font
                        
                    para_response = doc.add_paragraph(f"응답: {response.response[:500]}...")
                    for run in para_response.runs:
                        run.font.name = korean_font
                        
                    doc.add_paragraph("")  # 빈 줄
            
            # 파일 저장
            doc.save(output_path)
            return True
            
        except Exception as e:
            print(f"Word 내보내기 실패: {e}")
            return False

    def format_console_report(self, report: InvestmentReport, gpt_responses: List[GPTResponse] = None) -> str:
        """콘솔용 리포트 포맷팅"""
        lines = []

        # 헤더
        lines.append("=" * 80)
        lines.append(f"🦄 AI 스타트업 투자 평가 리포트: {report.company_info.name}")
        lines.append("=" * 80)

        # Executive Summary
        lines.append("\n📊 EXECUTIVE SUMMARY")
        lines.append("-" * 40)
        lines.append(f"종합 점수: {report.unicorn_score.total_score:.1f}/100 ({report.unicorn_score.grade}급)")
        lines.append(f"유니콘 확률: {report.unicorn_score.unicorn_probability:.1%}")
        lines.append(f"투자 추천: {report.recommendation.value}")
        lines.append(f"신뢰도: {report.confidence_level:.1%}")
        lines.append("")
        lines.append(report.executive_summary)

        # 영역별 점수
        lines.append("\n📈 영역별 점수카드")
        lines.append("-" * 40)
        for result in report.analysis_results:
            lines.append(f"{result.category:20} {result.score:5.1f}점 ({result.grade}급)")

        # 리스크 평가
        lines.append("\n⚠️ 리스크 평가")
        lines.append("-" * 40)
        for risk in report.risk_assessments:
            risk_emoji = {
                "낮음": "🟢",
                "보통": "🟡",
                "높음": "🟠",
                "매우 높음": "🔴"
            }
            emoji = risk_emoji.get(risk.risk_level.value, "⚪")
            lines.append(f"{emoji} {risk.category}: {risk.risk_level.value}")

        # 투자 권장사항
        lines.append(f"\n💰 투자 권장사항")
        lines.append("-" * 40)
        lines.append(report.investment_rationale)

        # 메타데이터
        lines.append(f"\n📋 평가 정보")
        lines.append("-" * 40)
        lines.append(f"평가 일시: {report.evaluation_date.strftime('%Y-%m-%d %H:%M:%S')}")
        lines.append(f"데이터 소스: {len(report.data_sources)}개")
        if report.limitations:
            lines.append(f"제한사항: {', '.join(report.limitations)}")

        # GPT 응답 로그 추가
        if gpt_responses:
            lines.append(self.format_gpt_responses_section(gpt_responses))

        lines.append("=" * 80)

        return "\n".join(lines)

    def format_json_report(self, report: InvestmentReport) -> Dict[str, Any]:
        """JSON 형식 리포트 포맷팅"""
        return {
            "company_info": {
                "name": report.company_info.name,
                "industry": report.company_info.industry,
                "founded_year": report.company_info.founded_year,
                "headquarters": report.company_info.headquarters,
                "description": report.company_info.description
            },
            "evaluation_summary": {
                "evaluation_date": report.evaluation_date.isoformat(),
                "total_score": report.unicorn_score.total_score,
                "grade": report.unicorn_score.grade,
                "unicorn_probability": report.unicorn_score.unicorn_probability,
                "recommendation": report.recommendation.value,
                "confidence_level": report.confidence_level
            },
            "analysis_results": [
                {
                    "category": result.category,
                    "score": result.score,
                    "grade": result.grade,
                    "summary": result.summary,
                    "strengths": result.key_strengths,
                    "weaknesses": result.key_weaknesses
                }
                for result in report.analysis_results
            ],
            "risk_assessments": [
                {
                    "category": risk.category,
                    "risk_level": risk.risk_level.value,
                    "description": risk.description,
                    "impact_score": risk.impact_score,
                    "probability": risk.probability,
                    "mitigation_strategies": risk.mitigation_strategies
                }
                for risk in report.risk_assessments
            ],
            "reports": {
                "executive_summary": report.executive_summary,
                "detailed_analysis": report.detailed_analysis,
                "investment_rationale": report.investment_rationale,
                "risk_summary": report.risk_summary
            },
            "metadata": {
                "data_sources": report.data_sources,
                "limitations": report.limitations,
                "score_breakdown": report.unicorn_score.score_breakdown
            }
        }

class ReportGenerationLayer:
    """리포트 생성 레이어 메인 클래스"""

    def __init__(self):
        self.report_generator = UnicornReportGenerator()
        self.formatter = ReportFormatter()

    def generate_report(
        self,
        company_info: CompanyInfo,
        unicorn_score: UnicornScore,
        analysis_results: List[AnalysisResult],
        risk_assessments: List[RiskAssessment],
        documents: List[DocumentChunk],
        external_results: List[ExternalSearchResult],
        gpt_responses: List[GPTResponse] = None
    ) -> InvestmentReport:
        """투자 평가 리포트 생성"""
        return self.report_generator.generate_investment_report(
            company_info=company_info,
            unicorn_score=unicorn_score,
            analysis_results=analysis_results,
            risk_assessments=risk_assessments,
            documents=documents,
            external_results=external_results
        )

    def format_report(self, report: InvestmentReport, format_type: str = "console", gpt_responses: List[GPTResponse] = None) -> str:
        """리포트 포맷팅"""
        if format_type == "console":
            return self.formatter.format_console_report(report, gpt_responses)
        elif format_type == "json":
            import json
            return json.dumps(self.formatter.format_json_report(report), ensure_ascii=False, indent=2)
        else:
            return self.formatter.format_console_report(report, gpt_responses)

    def export_report(self, report: InvestmentReport, gpt_responses: List[GPTResponse], output_path: str, format_type: str = "pdf") -> bool:
        """리포트를 파일로 내보내기"""
        if format_type == "pdf":
            return self.formatter.export_to_pdf(report, gpt_responses, output_path)
        elif format_type == "word":
            return self.formatter.export_to_word(report, gpt_responses, output_path)
        else:
            return False

def create_report_generation_layer() -> ReportGenerationLayer:
    """Report Generation Layer 생성자"""
    return ReportGenerationLayer()

def process_report_generation_layer(context: PipelineContext) -> PipelineContext:
    """Report Generation Layer 처리 함수"""
    report_layer = create_report_generation_layer()

    # 최종 투자 평가 리포트 생성
    investment_report = report_layer.generate_report(
        company_info=context.company_info,
        unicorn_score=context.unicorn_score,
        analysis_results=context.analysis_results,
        risk_assessments=context.risk_assessments,
        documents=context.retrieved_documents,
        external_results=context.external_search_results,
        gpt_responses=context.gpt_responses
    )

    context.final_report = investment_report

    # 처리 단계 기록
    context.processing_steps.append(
        f"REPORT_GENERATION_LAYER: 투자 평가 리포트 생성 완료 "
        f"(추천: {investment_report.recommendation.value})"
    )

    return context